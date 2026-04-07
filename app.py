import streamlit as st
import pandas as pd
import PyPDF2
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from io import BytesIO

# 1. Función para leer la estrategia y armar los nombres exactos
def obtener_matriz_estrategia():
    try:
        df = pd.read_csv('estrategia.csv')
        criterios = ""
        nombres_planteos = []
        for _, fila in df.iterrows():
            nombre = str(fila['Planteo']).strip()
            nombres_planteos.append(nombre)
            criterios += f"- {nombre}: Buscar '{fila['PalabrasClave']}'. Ejemplo: '{fila['Ejemplo']}'\n"
        return criterios, nombres_planteos
    except Exception as e:
        return f"Error al cargar matriz: {e}", []

# 2. Función MAESTRA para crear el Word con TABLAS REALES
def crear_word_profesional(texto_ia):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    lineas = texto_ia.split('\n')
    en_tabla = False
    tabla_word = None

    for linea in lineas:
        linea_limpia = linea.strip()
        if not linea_limpia: continue

        # Identificar si es una fila de tabla de la IA
        if linea_limpia.startswith('|') and linea_limpia.count('|') > 1:
            # Ignorar la línea de separación de Markdown (incluyendo los dos puntos)
            if set(linea_limpia.replace('|', '').replace('-', '').replace(':', '').replace(' ', '')) == set():
                continue
            
            # Extraer celdas
            celdas_texto = [c.strip() for c in linea_limpia.split('|') if c.strip()]
            
            if not en_tabla:
                # Crear tabla nueva
                tabla_word = doc.add_table(rows=1, cols=len(celdas_texto))
                tabla_word.style = 'Table Grid'
                hdr_cells = tabla_word.rows[0].cells
                for i, texto in enumerate(celdas_texto):
                    hdr_cells[i].text = texto
                en_tabla = True
            else:
                # Agregar fila a tabla existente
                row_cells = tabla_word.add_row().cells
                for i, texto in enumerate(celdas_texto):
                    if i < len(row_cells):
                        row_cells[i].text = texto
        else:
            en_tabla = False
            # Manejo de Títulos
            if linea_limpia.startswith('###'):
                doc.add_heading(linea_limpia.replace('###', '').strip(), level=2)
            elif linea_limpia.startswith('##'):
                doc.add_heading(linea_limpia.replace('##', '').strip(), level=1)
            elif linea_limpia.startswith('#'):
                doc.add_heading(linea_limpia.replace('#', '').strip(), level=0)
            else:
                doc.add_paragraph(linea_limpia.replace('**', ''))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# 3. Interfaz Streamlit
st.set_page_config(page_title="Analizador Pro v3", page_icon="⚖️")
st.title("⚖️ Analizador de planteos en Contestaciones de demandas de ANSES")

st.sidebar.header("Configuración")
api_key = st.sidebar.text_input("OpenAI API Key", type="password")

archivo = st.file_uploader("Sube la contestación de ANSES (PDF)", type="pdf")

if st.button("🚀 Iniciar Análisis Profesional"):
    if not api_key or not archivo:
        st.error("Falta la API Key o el archivo.")
    else:
        client = OpenAI(api_key=api_key)
        
        with st.spinner("Extrayendo datos e inyectando marcadores de página..."):
            lector = PyPDF2.PdfReader(archivo)
            texto_demanda = ""
            # CAMBIO CLAVE: Inyectamos el número de página real para que la IA lo vea
            for num_pag, pagina in enumerate(lector.pages, start=1):
                texto_demanda += f"\n\n[[[--- PÁGINA {num_pag} ---]]]\n\n"
                texto_demanda += pagina.extract_text() + "\n"
            
            # Limpieza de espacios extra pero manteniendo las marcas de página
            texto_demanda = " ".join(texto_demanda.split())
            
            criterios, nombres_planteos = obtener_matriz_estrategia()
            
            if not nombres_planteos:
                st.error("No se pudo cargar el archivo CSV. Verifica que 'estrategia.csv' esté en la carpeta.")
            else:
                # CAMBIO CLAVE: Armamos la plantilla exacta que la IA DEBE respetar
                esqueleto_tabla = "| Planteo | ¿Está? (SÍ/NO) | Evidencia Textual | Página | Certeza |\n"
                esqueleto_tabla += "| :--- | :--- | :--- | :--- | :--- |\n"
                for nombre in nombres_planteos:
                    esqueleto_tabla += f"| {nombre} | [SÍ o NO] | [Párrafo breve o '-'] | [Nro Pág o '-'] | [ALTA/MEDIA/BAJA o NULA] |\n"

                prompt_sistema = f"""Actúa como un Auditor de Juzgado Especialista en Seguridad Social. Eres estricto, meticuloso y no inventas información.

Tu tarea es analizar el documento adjunto (que contiene marcas como [[[--- PÁGINA X ---]]] para guiarte) y evaluar la presencia de ciertos planteos.

1. DATOS DE IDENTIFICACIÓN:
Extrae Carátula, Expediente y Juzgado.

2. CRITERIOS DE EVALUACIÓN:
Evalúa el texto buscando la existencia de estos planteos:
{criterios}

3. INSTRUCCIÓN CRÍTICA (FORMULARIO ESTRICTO):
Está TERMINANTEMENTE PROHIBIDO inventar planteos nuevos, cambiarles el nombre o alterar su orden. Debes COPIAR la siguiente plantilla exacta de tabla y REEMPLAZAR los campos entre corchetes [...] con la información extraída.

COPIA Y COMPLETA ESTA TABLA EXACTA:
{esqueleto_tabla}

REGLAS DE RELLENO:
- ¿Está?: Responde únicamente "SÍ" o "NO".
- Evidencia Textual: Si está, cita un fragmento breve (1 o 2 oraciones máximo). Si no está, pon "-".
- Página: Usa las marcas [[[--- PÁGINA X ---]]] del texto para indicar el número. Si no está, pon "-".
- Certeza: Si está, indica ALTA, MEDIA o BAJA. Si no está, indica NULA.

FORMATO DE SALIDA REQUERIDO:
# REPORTE DE ANÁLISIS LEGAL
## DATOS DEL EXPEDIENTE
- **Carátula:** [Nombre]
- **Expediente:** [Número]
- **Juzgado:** [Número]

## RESULTADO DEL MAPEO DE DEFENSA
(Pega aquí la tabla completada)"""

                try:
                    res = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "system", "content": prompt_sistema},
                            {"role": "user", "content": texto_demanda}
                        ],
                        temperature=0
                    )
                    
                    informe_final = res.choices[0].message.content
                    st.success("¡Análisis Terminado con éxito!")
                    st.markdown(informe_final)
                    
                    word_ready = crear_word_profesional(informe_final)
                    
                    st.download_button(
                        label="📥 Descargar Informe en Word",
                        data=word_ready,
                        file_name="Informe_Estrategia_ANSES.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Error detectado: {e}")
